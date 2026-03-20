"""Writer module for .docx (Microsoft Word) output."""

from html.parser import HTMLParser
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from mtd.models import Document
from mtd.themes.engine import Theme

# ---------------------------------------------------------------------------
# Measurement helpers
# ---------------------------------------------------------------------------


def _parse_cm(value: str) -> float:
    """Parse a measurement string like '2.54cm' to float cm value."""
    value = value.strip().lower()
    if value.endswith("cm"):
        return float(value[:-2])
    if value.endswith("in"):
        return float(value[:-2]) * 2.54
    return float(value)


def _parse_color(hex_color: str) -> RGBColor:
    """Parse a hex color string like '#2563eb' into an RGBColor."""
    hex_color = hex_color.lstrip("#")
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


def _hex_to_fill(hex_color: str) -> str:
    """Convert a hex color string to a 6-char uppercase string for XML fill."""
    return hex_color.lstrip("#").upper()


# ---------------------------------------------------------------------------
# Internal HTML -> DOCX converter
# ---------------------------------------------------------------------------

_HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
_BLOCK_TAGS = _HEADING_TAGS | {"p", "li", "blockquote", "pre", "hr", "div"}
_INLINE_BOLD = {"strong", "b"}
_INLINE_ITALIC = {"em", "i"}
_INLINE_CODE = {"code"}
_INLINE_STRIKE = {"del", "s"}
_INLINE_ANCHOR = {"a"}


def _add_shading_to_paragraph(paragraph, fill_color: str = "F2F2F2") -> None:
    """Apply a background shading to a paragraph via XML."""
    pPr = paragraph._p.get_or_add_pPr()  # noqa: N806
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    pPr.append(shd)


def _add_paragraph_border(
    paragraph, position: str = "bottom", color: str = "000000", size: str = "6", space: str = "1"
) -> None:
    """Add a border to a paragraph (used for <hr>)."""
    pPr = paragraph._p.get_or_add_pPr()  # noqa: N806
    pBdr = OxmlElement("w:pBdr")  # noqa: N806
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:color"), color)
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    pBdr.append(border)
    pPr.append(pBdr)


def _set_paragraph_indent(paragraph, left_twips: int = 720) -> None:
    """Set left indentation on a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()  # noqa: N806
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    pPr.append(ind)


def _add_page_number(run) -> None:
    """Insert a PAGE field into the given run element."""
    fldSimple = OxmlElement("w:fldSimple")  # noqa: N806
    fldSimple.set(qn("w:instr"), "PAGE")
    run._element.append(fldSimple)


def _add_header_footer_tab_stops(paragraph) -> None:
    """Add center and right tab stops to a header/footer paragraph."""
    pf = paragraph.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)


def _build_hf_paragraph(
    paragraph, left: str, center: str, right: str, date_str: str, title_str: str
) -> None:
    """Populate a header or footer paragraph with left/center/right content.

    Replaces {page} with a page number field, {date} with the date string,
    and {title} with the document title.
    """
    _add_header_footer_tab_stops(paragraph)

    def _add_part(text: str) -> None:
        """Add a text part, handling {page} specially."""
        parts = text.split("{page}")
        for i, part in enumerate(parts):
            part = part.replace("{date}", date_str).replace("{title}", title_str)
            if part:
                paragraph.add_run(part)
            if i < len(parts) - 1:
                # Insert page number field
                run = paragraph.add_run()
                _add_page_number(run)

    _add_part(left)
    if center or right:
        paragraph.add_run("\t")
        _add_part(center)
    if right:
        paragraph.add_run("\t")
        _add_part(right)


class _DocxHTMLParser(HTMLParser):
    """Walk an HTML string and populate a python-docx Document."""

    def __init__(
        self, docx: DocxDocument, titlepage_mode: bool = False, theme: Theme | None = None
    ) -> None:
        super().__init__(convert_charrefs=True)
        self._docx = docx
        self._titlepage_mode = titlepage_mode
        self._theme = theme or Theme()
        self._first_element = True

        # Current state
        self._tag_stack: list[tuple[str, dict]] = []  # (tag, attrs_dict)

        # Inline formatting flags (accumulated from stack)
        self._bold: int = 0
        self._italic: int = 0
        self._code_inline: int = 0
        self._strike: int = 0

        # Current paragraph being built; None means we need to create one
        self._current_para = None
        self._para_style: str | None = None  # deferred style for next para

        # Pre / code block tracking
        self._in_pre: int = 0

        # List tracking
        self._list_stack: list[str] = []  # "ul" or "ol"

        # Blockquote depth
        self._blockquote_depth: int = 0

        # Table accumulation: list of rows, each row is list of cell-text strings
        self._in_table: bool = False
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell_parts: list[str] = []
        self._in_th: bool = False

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _attrs_dict(self, attrs: list[tuple[str, str | None]]) -> dict:
        return {k: v for k, v in attrs}

    def _flush_cell(self) -> None:
        self._current_cell_parts = []

    def _commit_cell(self) -> None:
        self._current_row.append("".join(self._current_cell_parts))
        self._current_cell_parts = []

    def _commit_row(self) -> None:
        if self._current_row:
            self._table_rows.append(self._current_row)
        self._current_row = []

    def _ensure_paragraph(self, style: str = "Normal") -> None:
        """Return the current paragraph, creating one if necessary."""
        if self._current_para is None:
            self._current_para = self._docx.add_paragraph(style=style)

    def _new_paragraph(self, style: str = "Normal") -> None:
        """Finalise any open paragraph and start a fresh one."""
        self._current_para = self._docx.add_paragraph(style=style)

    def _apply_titlepage_style(self, tag: str) -> None:
        """Apply titlepage formatting to the current paragraph."""
        if self._current_para is None:
            return
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self._current_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        theme = self._theme
        if tag == "h1":
            size = Pt(theme.titlepage_title_size)
            bold = True
        elif tag == "h2":
            size = Pt(theme.titlepage_subtitle_size)
            bold = False
        else:
            size = Pt(theme.titlepage_info_size)
            bold = False
        for run in self._current_para.runs:
            run.font.size = size
            run.font.name = theme.font_heading
            if bold:
                run.bold = True
        # Store for runs added later
        self._current_para._mtd_tp_size = size
        self._current_para._mtd_tp_bold = bold

    def _add_run(self, text: str) -> None:
        """Append a formatted run to the current paragraph."""
        if not text:
            return
        self._ensure_paragraph()
        run = self._current_para.add_run(text)
        theme = self._theme
        if self._bold:
            run.bold = True
        if self._italic:
            run.italic = True
        if self._strike:
            run.font.strike = True
        if self._code_inline or self._in_pre:
            run.font.name = theme.font_code
            run.font.size = Pt(theme.code_size)
        else:
            # Apply body font and size for regular text
            run.font.name = theme.font_body
            run.font.size = Pt(theme.body_size)
            run.font.color.rgb = _parse_color(theme.color_text)
        # Apply titlepage font sizing if in titlepage mode
        if self._titlepage_mode and hasattr(self._current_para, "_mtd_tp_size"):
            run.font.size = self._current_para._mtd_tp_size
            run.font.name = theme.font_heading
            if getattr(self._current_para, "_mtd_tp_bold", False):
                run.bold = True
        return run

    # ------------------------------------------------------------------
    # Tag handlers
    # ------------------------------------------------------------------

    def handle_starttag(self, tag: str, attrs: list[tuple]) -> None:
        attrs_d = self._attrs_dict(attrs)
        self._tag_stack.append((tag, attrs_d))

        # ---- Table handling ----
        if tag == "table":
            self._in_table = True
            self._table_rows = []
            self._current_row = []
            self._current_cell_parts = []
            return
        if tag == "tr":
            self._current_row = []
            return
        if tag in ("th", "td"):
            self._in_th = tag == "th"
            self._current_cell_parts = []
            return

        # While inside a table cell, accumulate text only
        if self._in_table:
            return

        # ---- Block elements ----
        if tag in _HEADING_TAGS:
            level = int(tag[1])
            self._new_paragraph(style=f"Heading {level}")
            if self._titlepage_mode:
                self._current_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if self._first_element:
                    self._current_para.paragraph_format.space_before = Pt(72)
                    self._first_element = False
                # Store titlepage sizing info for runs
                theme = self._theme
                if tag == "h1":
                    self._current_para._mtd_tp_size = Pt(theme.titlepage_title_size)
                    self._current_para._mtd_tp_bold = True
                elif tag == "h2":
                    self._current_para._mtd_tp_size = Pt(theme.titlepage_subtitle_size)
                    self._current_para._mtd_tp_bold = False
                else:
                    self._current_para._mtd_tp_size = Pt(theme.titlepage_info_size)
                    self._current_para._mtd_tp_bold = False
            return

        if tag == "p":
            self._new_paragraph(style="Normal")
            if self._titlepage_mode:
                self._current_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if self._first_element:
                    self._current_para.paragraph_format.space_before = Pt(72)
                    self._first_element = False
                self._current_para._mtd_tp_size = Pt(self._theme.titlepage_info_size)
                self._current_para._mtd_tp_bold = False
            return

        if tag == "pre":
            self._in_pre += 1
            self._new_paragraph(style="Normal")
            if self._current_para is not None:
                fill = _hex_to_fill(self._theme.color_code_background)
                _add_shading_to_paragraph(self._current_para, fill)
            return

        if tag == "blockquote":
            self._blockquote_depth += 1
            self._new_paragraph(style="Normal")
            if self._current_para is not None:
                _set_paragraph_indent(self._current_para, 720 * self._blockquote_depth)
            return

        if tag == "ul":
            self._list_stack.append("ul")
            return

        if tag == "ol":
            self._list_stack.append("ol")
            return

        if tag == "li":
            if self._list_stack and self._list_stack[-1] == "ol":
                self._new_paragraph(style="List Number")
            else:
                self._new_paragraph(style="List Bullet")
            return

        if tag == "hr":
            p = self._docx.add_paragraph()
            _add_paragraph_border(p, position="bottom")
            self._current_para = None
            return

        # ---- Inline elements ----
        if tag in _INLINE_BOLD:
            self._bold += 1
            return

        if tag in _INLINE_ITALIC:
            self._italic += 1
            return

        if tag in _INLINE_CODE:
            if self._in_pre == 0:
                self._code_inline += 1
            return

        if tag in _INLINE_STRIKE:
            self._strike += 1
            return

        if tag == "a":
            # Hyperlinks: just render the inner text as-is (no real link)
            return

        if tag == "br":
            if self._current_para is not None:
                self._current_para.add_run().add_break()
            return

        if tag == "img":
            # Images: not supported; skip silently
            return

    def handle_endtag(self, tag: str) -> None:
        # Pop from stack if it matches
        for i in range(len(self._tag_stack) - 1, -1, -1):
            if self._tag_stack[i][0] == tag:
                self._tag_stack.pop(i)
                break

        # ---- Table handling ----
        if tag in ("td", "th"):
            self._commit_cell()
            return

        if tag == "tr":
            self._commit_row()
            return

        if tag == "table":
            self._in_table = False
            self._render_table()
            self._current_para = None
            return

        # ---- Block elements ----
        if tag in _HEADING_TAGS or tag == "p":
            self._current_para = None
            return

        if tag == "pre":
            self._in_pre = max(0, self._in_pre - 1)
            self._current_para = None
            return

        if tag == "blockquote":
            self._blockquote_depth = max(0, self._blockquote_depth - 1)
            self._current_para = None
            return

        if tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
            self._current_para = None
            return

        if tag == "li":
            self._current_para = None
            return

        # ---- Inline elements ----
        if tag in _INLINE_BOLD:
            self._bold = max(0, self._bold - 1)
            return

        if tag in _INLINE_ITALIC:
            self._italic = max(0, self._italic - 1)
            return

        if tag in _INLINE_CODE:
            if self._in_pre == 0:
                self._code_inline = max(0, self._code_inline - 1)
            return

        if tag in _INLINE_STRIKE:
            self._strike = max(0, self._strike - 1)
            return

    def handle_data(self, data: str) -> None:
        if self._in_table:
            # Accumulate text for current cell
            self._current_cell_parts.append(data)
            return

        # Outside a pre block, skip data that is pure whitespace and we have
        # no open paragraph (i.e. it is inter-block whitespace from the HTML).
        if self._in_pre == 0 and self._current_para is None and not data.strip():
            return

        self._add_run(data)

    # ------------------------------------------------------------------
    # Table renderer
    # ------------------------------------------------------------------

    def _render_table(self) -> None:
        if not self._table_rows:
            return
        rows = self._table_rows
        col_count = max(len(r) for r in rows)
        table = self._docx.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                if c_idx >= col_count:
                    break
                cell = row.cells[c_idx]
                cell.text = cell_text.strip()
                if r_idx == 0:
                    # Bold header row
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True


# ---------------------------------------------------------------------------
# Heading style application
# ---------------------------------------------------------------------------


def _apply_heading_styles(docx: DocxDocument, theme: Theme) -> None:
    """Apply theme heading styles to the document's built-in Heading styles."""
    heading_styles = {
        1: theme.h1,
        2: theme.h2,
        3: theme.h3,
        4: theme.h4,
        5: theme.h5,
        6: theme.h6,
    }
    for level, h_style in heading_styles.items():
        style_name = f"Heading {level}"
        try:
            style = docx.styles[style_name]
        except KeyError:
            continue
        font = style.font
        font.name = theme.font_heading
        font.size = Pt(h_style.size)
        font.bold = h_style.bold
        font.italic = h_style.italic
        if h_style.color:
            font.color.rgb = _parse_color(h_style.color)
        else:
            font.color.rgb = _parse_color(theme.color_primary)


# ---------------------------------------------------------------------------
# Header / Footer helpers
# ---------------------------------------------------------------------------


def _resolve_placeholders(text: str, date_str: str, title_str: str) -> str:
    """Replace {date} and {title} in text (not {page}, handled separately)."""
    return text.replace("{date}", date_str).replace("{title}", title_str)


def _setup_header(section, header_dict: dict, date_str: str, title_str: str) -> None:
    """Add header content to a DOCX section."""
    header = section.header
    # Clear default empty paragraph
    for para in list(header.paragraphs):
        para._element.getparent().remove(para._element)

    para = header.add_paragraph()
    left = _resolve_placeholders(header_dict.get("left", ""), date_str, title_str)
    center = _resolve_placeholders(header_dict.get("center", ""), date_str, title_str)
    right = header_dict.get("right", "")

    _build_hf_paragraph(para, left, center, right, date_str, title_str)


def _setup_footer(section, footer_dict: dict, date_str: str, title_str: str) -> None:
    """Add footer content to a DOCX section."""
    footer = section.footer
    # Clear default empty paragraph
    for para in list(footer.paragraphs):
        para._element.getparent().remove(para._element)

    para = footer.add_paragraph()
    left = _resolve_placeholders(footer_dict.get("left", ""), date_str, title_str)
    center = _resolve_placeholders(footer_dict.get("center", ""), date_str, title_str)
    right = footer_dict.get("right", "")

    _build_hf_paragraph(para, left, center, right, date_str, title_str)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def write_docx(document: Document, output: str | Path, theme: Theme | None = None) -> Path:
    """Convert a Document to DOCX format.

    Args:
        document: Parsed Document object.
        output: Output file path.
        theme: Optional Theme to apply. Defaults to Theme() if not provided.

    Returns:
        The path to the created DOCX file.
    """
    if theme is None:
        theme = Theme()

    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)

    docx = DocxDocument()

    # Apply heading styles from theme
    _apply_heading_styles(docx, theme)

    # Apply body text style defaults
    try:
        normal_style = docx.styles["Normal"]
        normal_style.font.name = theme.font_body
        normal_style.font.size = Pt(theme.body_size)
        normal_style.font.color.rgb = _parse_color(theme.color_text)
        from docx.oxml import OxmlElement as _OxmlElement
        from docx.oxml.ns import qn as _qn

        pPr = normal_style.element.get_or_add_pPr()  # noqa: N806
        spacing = _OxmlElement("w:spacing")
        spacing.set(_qn("w:line"), str(int(theme.line_spacing * 240)))
        spacing.set(_qn("w:lineRule"), "auto")
        pPr.append(spacing)
    except (KeyError, Exception):
        pass

    # Set core properties from metadata
    props = docx.core_properties
    if document.title:
        props.title = document.title
    if document.author:
        props.author = document.author
    if document.date:
        props.created = _parse_date(document.date)

    # Remove the default empty paragraph that python-docx creates
    for para in list(docx.paragraphs):
        para._element.getparent().remove(para._element)

    # Resolve date and title strings for placeholders
    date_str = str(document.date) if document.date else ""
    title_str = document.title or ""

    # Configure section for header/footer and page margins
    section = docx.sections[0]

    # Apply page margins from theme
    section.top_margin = Cm(_parse_cm(theme.margin_top))
    section.bottom_margin = Cm(_parse_cm(theme.margin_bottom))
    section.left_margin = Cm(_parse_cm(theme.margin_left))
    section.right_margin = Cm(_parse_cm(theme.margin_right))

    has_titlepage = document.titlepage is not None
    has_header = document.header is not None
    has_footer = document.footer is not None

    if has_titlepage and (has_header or has_footer):
        # First page should not show header/footer
        section.different_first_page_header_footer = True

    if has_header:
        _setup_header(section, document.header, date_str, title_str)

    if has_footer:
        _setup_footer(section, document.footer, date_str, title_str)

    # Render titlepage content first
    if has_titlepage:
        tp_parser = _DocxHTMLParser(docx, titlepage_mode=True, theme=theme)
        tp_parser.feed(document.titlepage)
        # Add a page break after the titlepage using a run-level break
        pb_para = docx.add_paragraph()
        pb_run = pb_para.add_run()
        pb_run.add_break(WD_BREAK.PAGE)

    parser = _DocxHTMLParser(docx, theme=theme)
    parser.feed(document.content)

    docx.save(str(output))
    return output


def _parse_date(date_str: str):
    """Attempt to parse a date string into a datetime, return as-is on failure."""
    from datetime import datetime

    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%d/%m/%Y", "%B %d, %Y"):
        try:
            return datetime.strptime(str(date_str), fmt)
        except (ValueError, TypeError):
            pass
    return None
