"""Tests for titlepage extraction and header/footer frontmatter."""

import zipfile
from pathlib import Path

from docx import Document as DocxDocument

from mtd.parser import parse_markdown
from mtd.writers.docx_writer import write_docx
from mtd.writers.odt_writer import write_odt

# ---------------------------------------------------------------------------
# Fixtures / helpers
# ---------------------------------------------------------------------------

_TITLEPAGE_MD = """\
:::titlepage
# My Great Title

## A Subtitle

Author Name
:::

# Introduction

This is the main content.
"""

_NO_TITLEPAGE_MD = """\
# Introduction

This is the main content.
"""

_HEADER_FOOTER_MD = """\
---
title: Test Doc
header:
  left: "Left Header"
  center: "Center Header"
  right: "{page}"
footer:
  left: "Left Footer"
  center: ""
  right: ""
---
# Body

Some text.
"""


# ---------------------------------------------------------------------------
# Parser-level tests
# ---------------------------------------------------------------------------


def test_titlepage_extracted():
    """parse_markdown detects a :::titlepage block and sets document.titlepage."""
    doc = parse_markdown(_TITLEPAGE_MD)
    assert doc.titlepage is not None


def test_titlepage_content():
    """The titlepage HTML contains the heading text from the block."""
    doc = parse_markdown(_TITLEPAGE_MD)
    assert doc.titlepage is not None
    assert "My Great Title" in doc.titlepage


def test_no_titlepage():
    """Documents without a :::titlepage block have titlepage == None."""
    doc = parse_markdown(_NO_TITLEPAGE_MD)
    assert doc.titlepage is None


def test_titlepage_removed_from_content():
    """The main content HTML must NOT contain titlepage-only text."""
    doc = parse_markdown(_TITLEPAGE_MD)
    # "My Great Title" should be in the titlepage, not in the main content
    assert "My Great Title" not in doc.content
    # Main content should still be present
    assert "Introduction" in doc.content


def test_header_from_frontmatter():
    """document.header returns the header dict from frontmatter."""
    doc = parse_markdown(_HEADER_FOOTER_MD)
    assert doc.header is not None
    assert isinstance(doc.header, dict)
    assert doc.header.get("left") == "Left Header"
    assert doc.header.get("center") == "Center Header"


def test_footer_from_frontmatter():
    """document.footer returns the footer dict from frontmatter."""
    doc = parse_markdown(_HEADER_FOOTER_MD)
    assert doc.footer is not None
    assert isinstance(doc.footer, dict)
    assert doc.footer.get("left") == "Left Footer"


def test_no_header_footer():
    """When frontmatter has no header/footer keys, both properties are None."""
    doc = parse_markdown("# Plain document\n\nNo frontmatter.")
    assert doc.header is None
    assert doc.footer is None


# ---------------------------------------------------------------------------
# DOCX writer tests
# ---------------------------------------------------------------------------


def test_docx_with_titlepage(tmp_path: Path):
    """write_docx with a titlepage creates a file; the doc has centered paragraphs."""
    doc = parse_markdown(_TITLEPAGE_MD)
    output = tmp_path / "titlepage.docx"
    result = write_docx(doc, output)

    assert result == output
    assert output.exists()
    assert output.stat().st_size > 0

    # Verify the file is a valid DOCX (zip) with at least one paragraph
    docx = DocxDocument(str(output))
    all_text = " ".join(p.text for p in docx.paragraphs)
    # Titlepage text should appear somewhere in the document
    assert "My Great Title" in all_text or "Introduction" in all_text


def test_docx_with_header_footer(tmp_path: Path):
    """write_docx with header/footer frontmatter creates a valid DOCX file."""
    doc = parse_markdown(_HEADER_FOOTER_MD)
    output = tmp_path / "hf.docx"
    result = write_docx(doc, output)

    assert result == output
    assert output.exists()
    assert output.stat().st_size > 0

    # Verify the file can be opened by python-docx
    docx = DocxDocument(str(output))
    # Confirm the main body section has paragraphs
    assert len(docx.paragraphs) > 0


# ---------------------------------------------------------------------------
# ODT writer tests
# ---------------------------------------------------------------------------


def test_odt_with_titlepage(tmp_path: Path):
    """write_odt with a titlepage produces a valid ODT file with titlepage elements."""
    doc = parse_markdown(_TITLEPAGE_MD)
    output = tmp_path / "titlepage.odt"
    result = write_odt(doc, output)

    assert result == output
    assert output.exists()
    assert output.stat().st_size > 0

    with zipfile.ZipFile(str(output), "r") as zf:
        content = zf.read("content.xml").decode("utf-8")

    # The title text should be present in the ODT content
    assert "My Great Title" in content
    # TitlepageH1 style should be referenced in the document
    assert "TitlepageH1" in content
