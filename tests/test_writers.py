"""Integration tests for the DOCX and ODT writers."""

import zipfile
from pathlib import Path

from docx import Document as DocxDocument

from mtd.models import Document
from mtd.parser import parse_markdown
from mtd.writers.docx_writer import write_docx
from mtd.writers.odt_writer import write_odt

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

EXAMPLES_DIR = Path(__file__).parent.parent / "examples"
TEMPLATE_PATH = EXAMPLES_DIR / "template.md"

# ---------------------------------------------------------------------------
# Shared test documents
# ---------------------------------------------------------------------------

_HEADING_MD = "# Test Heading\n\nSome paragraph text."

_BOLD_MD = "**bold text** in a paragraph."

_ITALIC_MD = "*italic text* in a paragraph."

_CODE_BLOCK_MD = "```\nprint('hello world')\n```"

_TABLE_MD = (
    "| Column A | Column B |\n"
    "|----------|----------|\n"
    "| cell 1   | cell 2   |\n"
    "| cell 3   | cell 4   |\n"
)

_LIST_MD = "- First item\n- Second item\n- Third item"

_METADATA_MD = "---\ntitle: Test Document Title\nauthor: Test Author\n---\n# Content\n"


def _doc(md: str) -> Document:
    """Parse a Markdown string into a Document."""
    return parse_markdown(md)


# ===========================================================================
# DOCX tests
# ===========================================================================


def test_docx_creates_file(tmp_path: Path) -> None:
    """write_docx produces a file that exists and is non-empty."""
    output = tmp_path / "output.docx"
    result = write_docx(_doc(_HEADING_MD), output)
    assert result == output
    assert output.exists()
    assert output.stat().st_size > 0


def test_docx_contains_heading(tmp_path: Path) -> None:
    """Output DOCX contains the heading text."""
    output = tmp_path / "heading.docx"
    write_docx(_doc(_HEADING_MD), output)

    docx = DocxDocument(str(output))
    heading_texts = [p.text for p in docx.paragraphs if p.style.name.startswith("Heading")]
    assert any("Test Heading" in t for t in heading_texts)


def test_docx_contains_paragraph(tmp_path: Path) -> None:
    """Output DOCX contains paragraph text."""
    output = tmp_path / "paragraph.docx"
    write_docx(_doc(_HEADING_MD), output)

    docx = DocxDocument(str(output))
    all_text = " ".join(p.text for p in docx.paragraphs)
    assert "Some paragraph text" in all_text


def test_docx_bold_text(tmp_path: Path) -> None:
    """Output DOCX has at least one bold run."""
    output = tmp_path / "bold.docx"
    write_docx(_doc(_BOLD_MD), output)

    docx = DocxDocument(str(output))
    bold_runs = [run for para in docx.paragraphs for run in para.runs if run.bold]
    assert bold_runs, "Expected at least one bold run in the DOCX output"
    assert any("bold text" in run.text for run in bold_runs)


def test_docx_italic_text(tmp_path: Path) -> None:
    """Output DOCX has at least one italic run."""
    output = tmp_path / "italic.docx"
    write_docx(_doc(_ITALIC_MD), output)

    docx = DocxDocument(str(output))
    italic_runs = [run for para in docx.paragraphs for run in para.runs if run.italic]
    assert italic_runs, "Expected at least one italic run in the DOCX output"
    assert any("italic text" in run.text for run in italic_runs)


def test_docx_code_block(tmp_path: Path) -> None:
    """Output DOCX has runs using a monospace font (Courier New) for code blocks."""
    output = tmp_path / "code.docx"
    write_docx(_doc(_CODE_BLOCK_MD), output)

    docx = DocxDocument(str(output))
    mono_runs = [
        run for para in docx.paragraphs for run in para.runs if run.font.name == "Courier New"
    ]
    assert mono_runs, "Expected at least one Courier New run for code block"


def test_docx_table(tmp_path: Path) -> None:
    """Output DOCX contains a table with the expected data."""
    output = tmp_path / "table.docx"
    write_docx(_doc(_TABLE_MD), output)

    docx = DocxDocument(str(output))
    assert docx.tables, "Expected at least one table in the DOCX output"

    # Check cell content
    table = docx.tables[0]
    all_cell_text = [cell.text for row in table.rows for cell in row.cells]
    assert any("Column A" in t for t in all_cell_text)
    assert any("cell 1" in t for t in all_cell_text)


def test_docx_list(tmp_path: Path) -> None:
    """Output DOCX contains list items (List Bullet style paragraphs)."""
    output = tmp_path / "list.docx"
    write_docx(_doc(_LIST_MD), output)

    docx = DocxDocument(str(output))
    list_paras = [p for p in docx.paragraphs if "List" in p.style.name]
    assert list_paras, "Expected paragraphs with a List style in the DOCX output"
    all_text = " ".join(p.text for p in list_paras)
    assert "First item" in all_text
    assert "Second item" in all_text
    assert "Third item" in all_text


def test_docx_metadata_title(tmp_path: Path) -> None:
    """core_properties.title is set from frontmatter."""
    output = tmp_path / "meta.docx"
    write_docx(_doc(_METADATA_MD), output)

    docx = DocxDocument(str(output))
    assert docx.core_properties.title == "Test Document Title"


# ===========================================================================
# ODT tests
# ===========================================================================


def test_odt_creates_file(tmp_path: Path) -> None:
    """write_odt produces a file that exists and is non-empty."""
    output = tmp_path / "output.odt"
    result = write_odt(_doc(_HEADING_MD), output)
    assert result == output
    assert output.exists()
    assert output.stat().st_size > 0


def test_odt_valid_zip(tmp_path: Path) -> None:
    """Output ODT file is a valid ZIP archive."""
    output = tmp_path / "valid.odt"
    write_odt(_doc(_HEADING_MD), output)

    assert zipfile.is_zipfile(str(output)), "ODT file should be a valid ZIP archive"


def test_odt_contains_content_xml(tmp_path: Path) -> None:
    """ODT ZIP archive contains a content.xml entry."""
    output = tmp_path / "content_xml.odt"
    write_odt(_doc(_HEADING_MD), output)

    with zipfile.ZipFile(str(output), "r") as zf:
        names = zf.namelist()
    assert "content.xml" in names, f"Expected content.xml in ODT archive, got: {names}"


def test_odt_has_headings(tmp_path: Path) -> None:
    """content.xml contains ODF heading elements."""
    output = tmp_path / "headings.odt"
    write_odt(_doc(_HEADING_MD), output)

    with zipfile.ZipFile(str(output), "r") as zf:
        content = zf.read("content.xml").decode("utf-8")

    # ODF headings are expressed as <text:h ...>
    assert "text:h" in content, "Expected text:h element in content.xml"
    assert "Test Heading" in content


def test_odt_has_paragraphs(tmp_path: Path) -> None:
    """content.xml contains paragraph text."""
    output = tmp_path / "paragraphs.odt"
    write_odt(_doc(_HEADING_MD), output)

    with zipfile.ZipFile(str(output), "r") as zf:
        content = zf.read("content.xml").decode("utf-8")

    # ODF paragraphs are <text:p ...>
    assert "text:p" in content, "Expected text:p element in content.xml"
    assert "Some paragraph text" in content


def test_odt_has_table(tmp_path: Path) -> None:
    """content.xml contains ODF table elements."""
    output = tmp_path / "table.odt"
    write_odt(_doc(_TABLE_MD), output)

    with zipfile.ZipFile(str(output), "r") as zf:
        content = zf.read("content.xml").decode("utf-8")

    # ODF tables are <table:table ...>
    assert "table:table" in content, "Expected table:table element in content.xml"
    assert "Column A" in content
    assert "cell 1" in content


# ===========================================================================
# End-to-end tests
# ===========================================================================


def test_e2e_template_to_docx(tmp_path: Path) -> None:
    """Parse examples/template.md and write to DOCX without errors."""
    output = tmp_path / "template.docx"
    doc = parse_markdown(TEMPLATE_PATH)
    result = write_docx(doc, output)

    assert result == output
    assert output.exists()
    assert output.stat().st_size > 0

    # Verify basic document structure via python-docx
    docx = DocxDocument(str(output))

    # Metadata should be set
    assert docx.core_properties.title == "Document Title"

    # Headings must exist
    heading_paras = [p for p in docx.paragraphs if p.style.name.startswith("Heading")]
    assert heading_paras, "Expected heading paragraphs in the template DOCX"

    # Table must exist
    assert docx.tables, "Expected at least one table in the template DOCX"

    # List items must exist
    list_paras = [p for p in docx.paragraphs if "List" in p.style.name]
    assert list_paras, "Expected list paragraphs in the template DOCX"


def test_e2e_template_to_odt(tmp_path: Path) -> None:
    """Parse examples/template.md and write to ODT without errors."""
    output = tmp_path / "template.odt"
    doc = parse_markdown(TEMPLATE_PATH)
    result = write_odt(doc, output)

    assert result == output
    assert output.exists()
    assert output.stat().st_size > 0

    # Verify it is a valid ZIP and has required entries
    assert zipfile.is_zipfile(str(output))
    with zipfile.ZipFile(str(output), "r") as zf:
        names = zf.namelist()
        assert "content.xml" in names

        content = zf.read("content.xml").decode("utf-8")

    # Heading elements present
    assert "text:h" in content

    # Table elements present
    assert "table:table" in content

    # Known content from the template
    assert "Heading 1" in content or "Document Title" in content or "text:h" in content
